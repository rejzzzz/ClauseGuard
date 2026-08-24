# Comprehensive edge-case unit tests for redlining engine and Redliner Agent.
from pathlib import Path
from docx import Document

from backend.redlining.diff_utils import compute_word_diff, render_diff_html
from backend.redlining.docx_redline_engine import apply_tracked_redlines
from backend.agents.auditor.verdict_schema import ClauseVerdict, VerdictEnum, SeverityEnum
from backend.agents.redliner.agent import RedlinerAgent
from backend.agents.redliner.edit_schema import RedlineActionEnum

def test_diff_utils_edge_cases():
    # Identical text
    diff_same = compute_word_diff("Exact match text", "Exact match text")
    assert all(op == "equal" for op, _ in diff_same)
    
    # Completely replaced text
    html_diff = render_diff_html("Old Text", "New Replacement")
    assert "<del>Old Text</del>" in html_diff
    assert "<ins>New Replacement</ins>" in html_diff
    
    # Empty inputs
    assert render_diff_html("", "") == ""

def test_apply_tracked_redlines_delete_only(tmp_path: Path):
    input_docx = tmp_path / "delete_test.docx"
    output_docx = tmp_path / "delete_out.docx"
    
    doc = Document()
    doc.add_paragraph("Prohibited penalty clause.")
    doc.save(str(input_docx))
    
    edits = [{
        "original_text": "Prohibited penalty clause.",
        "proposed_text": "",
        "action": "DELETE"
    }]
    
    apply_tracked_redlines(input_docx, output_docx, edits)
    res_doc = Document(str(output_docx))
    xml_content = res_doc.paragraphs[0]._p.xml
    assert "w:del" in xml_content
    assert "w:ins" not in xml_content

def test_apply_tracked_redlines_insert_only(tmp_path: Path):
    input_docx = tmp_path / "insert_test.docx"
    output_docx = tmp_path / "insert_out.docx"
    
    doc = Document()
    doc.add_paragraph("Existing clause.")
    doc.save(str(input_docx))
    
    edits = [{
        "original_text": "Existing clause.",
        "proposed_text": "Inserted mandatory protection.",
        "action": "INSERT"
    }]
    
    apply_tracked_redlines(input_docx, output_docx, edits)
    res_doc = Document(str(output_docx))
    xml_content = res_doc.paragraphs[0]._p.xml
    assert "w:ins" in xml_content

def test_apply_tracked_redlines_unmatched_text(tmp_path: Path):
    input_docx = tmp_path / "unmatched_test.docx"
    output_docx = tmp_path / "unmatched_out.docx"
    
    doc = Document()
    doc.add_paragraph("Original untouched paragraph.")
    doc.save(str(input_docx))
    
    edits = [{
        "original_text": "Non existent text",
        "proposed_text": "Some text",
        "action": "REPLACE"
    }]
    
    apply_tracked_redlines(input_docx, output_docx, edits)
    res_doc = Document(str(output_docx))
    assert res_doc.paragraphs[0].text == "Original untouched paragraph."

def test_redliner_agent_missing_clause_verdict():
    agent = RedlinerAgent()
    verdict = ClauseVerdict(
        clause_id="c_missing",
        heading_title="Data Privacy",
        verdict=VerdictEnum.MISSING_CLAUSE,
        severity=SeverityEnum.HIGH,
        rationale="Missing GDPR data processing addendum."
    )
    edit = agent.generate_edit_for_verdict(verdict, "Data Privacy Section")
    assert edit is not None
    assert edit.action == RedlineActionEnum.INSERT
