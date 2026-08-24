# Unit tests for redlining OOXML mutation engine and diff utilities.
from pathlib import Path
from docx import Document
from backend.redlining.diff_utils import compute_word_diff, render_diff_html
from backend.redlining.clause_bank import ClauseBank
from backend.redlining.docx_redline_engine import apply_tracked_redlines

def test_compute_word_diff():
    diffs = compute_word_diff("uncapped liability", "capped at 1x fees")
    assert len(diffs) > 0
    ops = [op for op, _ in diffs]
    assert "delete" in ops or "insert" in ops

def test_render_diff_html():
    html = render_diff_html("Original text", "Proposed text")
    assert "<del>" in html or "<ins>" in html

def test_clause_bank_lookup():
    bank = ClauseBank(default_playbook="sample_vendor_msa")
    content = bank.get_rule_language("sample_vendor_msa_rule_1")
    assert content is not None
    assert len(content) > 0

def test_apply_tracked_redlines(tmp_path: Path):
    input_docx = tmp_path / "input.docx"
    output_docx = tmp_path / "output.docx"
    
    doc = Document()
    doc.add_paragraph("Vendor liability shall be uncapped.")
    doc.save(str(input_docx))
    
    edits = [{
        "original_text": "Vendor liability shall be uncapped.",
        "proposed_text": "Vendor liability shall be capped at 1x annual fees.",
        "action": "REPLACE"
    }]
    
    res = apply_tracked_redlines(input_docx, output_docx, edits)
    assert res.exists()
    
    # Reload generated docx and assert XML elements w:del and w:ins exist
    res_doc = Document(str(res))
    xml_str = res_doc.paragraphs[0]._p.xml
    assert "w:del" in xml_str
    assert "w:ins" in xml_str
