# End-to-end integration test for document auditing, redlining, and OOXML tracked-changes mutation.
from pathlib import Path
from docx import Document

from backend.ingestion.pipeline import ingest_contract
from backend.agents.auditor.agent import AuditorAgent
from backend.agents.redliner.agent import RedlinerAgent
from backend.redlining.docx_redline_engine import apply_tracked_redlines

def test_full_redlining_pipeline(tmp_path: Path):
    input_docx = tmp_path / "sample_vendor_contract.docx"
    redlined_docx = tmp_path / "redlined_vendor_contract.docx"
    
    # 1. Create sample contract with deviation
    doc = Document()
    doc.add_heading("Vendor Agreement", level=1)
    doc.add_heading("Section 1: Limitation of Liability", level=2)
    doc.add_paragraph("Vendor liability shall be uncapped for any operational losses.")
    doc.save(str(input_docx))
    
    # 2. Ingest document
    chunks = ingest_contract(input_docx)
    assert len(chunks) == 1
    
    # 3. Audit document
    auditor = AuditorAgent()
    audit_report = auditor.audit_contract(chunks, playbook_name="sample_vendor_msa", contract_name="sample_vendor_contract.docx")
    assert audit_report.total_clauses == 1
    assert audit_report.verdicts[0].verdict.value == "DEVIATION"
    
    # 4. Generate redline package
    redliner = RedlinerAgent()
    redline_package = redliner.generate_redline_package(audit_report, chunks)
    assert redline_package.total_edits == 1
    
    # 5. Apply OOXML tracked changes mutation
    edits_dict = [e.model_dump() for e in redline_package.edits]
    res_path = apply_tracked_redlines(input_docx, redlined_docx, edits_dict)
    
    assert res_path.exists()
    
    # Reload docx and verify native Track Changes XML tags
    res_doc = Document(str(res_path))
    xml_content = "".join([p._p.xml for p in res_doc.paragraphs])
    assert "w:del" in xml_content
    assert "w:ins" in xml_content
