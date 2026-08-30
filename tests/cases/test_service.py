# Unit tests for case document ingestion service.
from pathlib import Path
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from docx import Document

from backend.db.base import Base
from backend.db.repository import (
    db_create_case,
    db_add_case_document,
    db_get_document_chunks,
    db_get_case_document,
    db_list_timeline_events
)
from backend.cases.service import ingest_case_document
from backend.ingestion.embedder import BedrockEmbedder


@pytest.fixture
def db():
    engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
    Base.metadata.create_all(bind=engine)
    TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    doc_path = tmp_path / "sample_service_contract.docx"
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph("This agreement is executed between Vendor and Client.")
    doc.add_heading("Section 1. Confidentiality", level=2)
    doc.add_paragraph("Recipient shall strictly protect all proprietary data.")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def sample_docx_with_dates(tmp_path: Path) -> Path:
    doc_path = tmp_path / "sample_dated_contract.docx"
    doc = Document()
    doc.add_heading("Loan Agreement", level=1)
    doc.add_paragraph("On 15 January 2022, Borrower received the loan principal.")
    doc.add_heading("Section 2. Repayment Schedule", level=2)
    doc.add_paragraph("On 2023-05-10, first installment was deposited to Lender.")
    doc.save(str(doc_path))
    return doc_path


def test_ingest_case_document_success(db, sample_docx):
    case = db_create_case(db, title="Service Ingestion Matter")
    doc = db_add_case_document(
        db=db,
        case_id=case.id,
        filename="sample_service_contract.docx",
        file_type="docx",
        file_path=str(sample_docx)
    )

    embedder = BedrockEmbedder()
    updated_doc = ingest_case_document(
        db=db,
        case_id=case.id,
        doc_id=doc.id,
        file_path=sample_docx,
        embedder=embedder,
        extract_timeline=False
    )

    assert updated_doc.ingestion_status == "COMPLETED"
    assert updated_doc.chunk_count > 0

    chunks = db_get_document_chunks(db, doc.id)
    assert len(chunks) == updated_doc.chunk_count
    assert any("Confidentiality" in c.text for c in chunks)


def test_ingest_case_document_with_auto_timeline(db, sample_docx_with_dates):
    case = db_create_case(db, title="Dated Ingestion Matter")
    doc = db_add_case_document(
        db=db,
        case_id=case.id,
        filename="sample_dated_contract.docx",
        file_type="docx",
        file_path=str(sample_docx_with_dates)
    )

    embedder = BedrockEmbedder()
    updated_doc = ingest_case_document(
        db=db,
        case_id=case.id,
        doc_id=doc.id,
        file_path=sample_docx_with_dates,
        embedder=embedder,
        extract_timeline=True
    )

    assert updated_doc.ingestion_status == "COMPLETED"
    assert updated_doc.chunk_count > 0

    # Verify timeline events were auto-extracted
    events = db_list_timeline_events(db=db, case_id=case.id)
    assert len(events) >= 1
    assert any("15 January 2022" in (e.event_date_raw or "") or "2023-05-10" in (e.event_date_raw or "") for e in events)
