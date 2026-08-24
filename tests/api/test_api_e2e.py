# End-to-end integration tests verifying full REST API lifecycle with real multi-agent pipeline.
from pathlib import Path
from docx import Document
from fastapi.testclient import TestClient

from backend.api.main import app

client = TestClient(app)


def test_e2e_api_full_contract_review_lifecycle(tmp_path: Path):
    # 1. Create real test .docx file
    docx_path = tmp_path / "e2e_vendor_agreement.docx"
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_heading("Section 1: Limitation of Liability", level=2)
    doc.add_paragraph("Vendor liability shall be uncapped for any operational losses.")
    doc.add_heading("Section 2: Governing Law", level=2)
    doc.add_paragraph("This Agreement shall be governed by Delaware law.")
    doc.save(str(docx_path))

    # 2. Upload Contract
    with open(docx_path, "rb") as f:
        files = {"file": ("e2e_vendor_agreement.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        data = {"playbook_name": "sample_vendor_msa"}
        upload_res = client.post("/api/sessions/upload", files=files, data=data)

    assert upload_res.status_code == 201
    upload_data = upload_res.json()
    session_id = upload_data["session_id"]
    assert session_id is not None
    assert upload_data["status"] == "UNINITIALIZED"

    # 3. Trigger Audit Pipeline
    audit_res = client.post(f"/api/sessions/{session_id}/audit")
    assert audit_res.status_code == 200
    audit_data = audit_res.json()
    assert audit_data["current_state"] == "AWAITING_HUMAN"
    assert audit_data["audit_report"] is not None
    assert audit_data["critic_report"] is not None
    assert audit_data["redline_package"] is not None

    # 4. Fetch Verdicts
    verdicts_res = client.get(f"/api/sessions/{session_id}/verdicts")
    assert verdicts_res.status_code == 200
    verdicts_data = verdicts_res.json()
    assert verdicts_data["total_clauses"] == 2
    assert len(verdicts_data["verdicts"]) == 2
    assert len(verdicts_data["edits"]) >= 1

    # 5. Apply Human HITL Review Decision (EDIT target clause)
    target_edit = verdicts_data["edits"][0]
    target_clause_id = target_edit["clause_id"]
    hitl_body = {
        "decisions": [
            {
                "clause_id": target_clause_id,
                "action": "EDIT",
                "custom_text": "Vendor liability shall be capped at $500,000 for any operational losses."
            }
        ]
    }
    hitl_res = client.post(f"/api/sessions/{session_id}/hitl", json=hitl_body)
    assert hitl_res.status_code == 200
    hitl_data = hitl_res.json()
    assert hitl_data["current_state"] == "AWAITING_HUMAN"

    # 6. Finalize Redline Document
    redline_res = client.post(f"/api/sessions/{session_id}/redline")
    assert redline_res.status_code == 200
    redline_data = redline_res.json()
    assert redline_data["current_state"] == "FINALIZED"
    assert "download" in redline_data["download_url"]

    # 7. Download Redlined Document
    dl_res = client.get(f"/api/sessions/{session_id}/download")
    assert dl_res.status_code == 200
    assert len(dl_res.content) > 0
    assert dl_res.headers.get("content-type") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # 8. Export Audit Report
    report_res = client.get(f"/api/sessions/{session_id}/report")
    assert report_res.status_code == 200
    report_data = report_res.json()
    assert report_data["session_id"] == session_id
    assert report_data["audit_report"] is not None
    assert len(report_data["history"]) >= 5
