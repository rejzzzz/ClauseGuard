# Integration tests for verdict inspection, HITL decision recording, and redline document download.
from pathlib import Path
import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.api.main import app
from backend.api.session_manager import session_manager
from backend.agents.orchestrator.state_machine import SessionContext, SessionStateEnum
from backend.agents.auditor.verdict_schema import ContractAuditReport, ClauseVerdict, VerdictEnum, SeverityEnum
from backend.agents.redliner.edit_schema import RedlinePackage, EditInstruction, RedlineActionEnum

client = TestClient(app)


def setup_awaiting_human_session(session_id="test_review_s1"):
    ctx = session_manager.create_session(
        contract_path="/path/to/contract.docx",
        contract_name="review_contract",
        session_id=session_id
    )
    ctx.current_state = SessionStateEnum.AWAITING_HUMAN
    ctx.audit_report = ContractAuditReport(
        contract_name="review_contract",
        playbook_name="sample_vendor_msa",
        total_clauses=1,
        verdicts=[
            ClauseVerdict(
                clause_id="c1",
                heading_title="Indemnity",
                verdict=VerdictEnum.DEVIATION,
                severity=SeverityEnum.HIGH,
                rationale="Unlimited liability clause",
                playbook_citation_ids=["rule_1"]
            )
        ],
        overall_risk_level=SeverityEnum.HIGH
    )
    ctx.redline_package = RedlinePackage(
        contract_name="review_contract",
        total_edits=1,
        edits=[
            EditInstruction(
                clause_id="c1",
                heading_title="Indemnity",
                action=RedlineActionEnum.REPLACE,
                original_text="Supplier agrees to unlimited liability",
                proposed_text="Supplier liability capped at 1x annual fees",
                comment_text="Standard cap requirement"
            )
        ]
    )
    session_manager.update_session(ctx)
    return ctx


def test_get_session_verdicts():
    setup_awaiting_human_session("test_verdicts_s1")
    response = client.get("/api/sessions/test_verdicts_s1/verdicts")
    assert response.status_code == 200
    data = response.json()
    assert data["session_id"] == "test_verdicts_s1"
    assert len(data["verdicts"]) == 1
    assert data["verdicts"][0]["clause_id"] == "c1"
    assert data["overall_risk_level"] == "HIGH"


def test_submit_hitl_decisions():
    setup_awaiting_human_session("test_hitl_s1")
    body = {
        "decisions": [
            {
                "clause_id": "c1",
                "action": "EDIT",
                "custom_text": "Supplier liability capped at 2x annual fees"
            }
        ]
    }
    response = client.post("/api/sessions/test_hitl_s1/hitl", json=body)
    assert response.status_code == 200
    data = response.json()
    assert data["session_id"] == "test_hitl_s1"
    assert data["updated_edits_count"] == 1
    
    # Verify updated edit in session store
    ctx = session_manager.get_session("test_hitl_s1")
    assert ctx.redline_package.edits[0].proposed_text == "Supplier liability capped at 2x annual fees"


def test_submit_hitl_invalid_state():
    ctx = session_manager.create_session(
        contract_path="/path/to/c.docx",
        contract_name="uninit_contract",
        session_id="test_uninit_s1"
    )
    body = {"decisions": [{"clause_id": "c1", "action": "APPROVE"}]}
    response = client.post("/api/sessions/test_uninit_s1/hitl", json=body)
    assert response.status_code == 400
    assert "must be in 'AWAITING_HUMAN' state" in response.json()["detail"]


def test_finalize_redline_and_download(tmp_path):
    ctx = setup_awaiting_human_session("test_finalize_s1")
    
    # Create valid docx contract file
    dummy_contract = tmp_path / "review_contract.docx"
    doc = Document()
    doc.add_heading("Indemnity", level=1)
    doc.add_paragraph("Supplier agrees to unlimited liability")
    doc.save(str(dummy_contract))
    
    ctx.contract_path = str(dummy_contract)
    session_manager.update_session(ctx)
    
    # Finalize redline
    response = client.post("/api/sessions/test_finalize_s1/redline")
    assert response.status_code == 200
    res_data = response.json()
    assert res_data["session_id"] == "test_finalize_s1"
    assert res_data["current_state"] == "FINALIZED"
    assert "download" in res_data["download_url"]
    
    # Test download
    dl_response = client.get("/api/sessions/test_finalize_s1/download")
    assert dl_response.status_code == 200
    assert dl_response.headers.get("content-type") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
