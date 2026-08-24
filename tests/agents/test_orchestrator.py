# Unit tests for Orchestrator Agent state machine and delegation.
import pytest
from pathlib import Path
from docx import Document

from backend.agents.orchestrator.state_machine import (
    AuditStateMachine,
    SessionContext,
    SessionStateEnum,
    HumanDecisionEnum,
    ClauseHumanDecision
)
from backend.agents.orchestrator.agent import OrchestratorAgent
from backend.agents.orchestrator.tools import (
    dispatch_to_ingestion,
    dispatch_to_auditor,
    dispatch_to_critic,
    dispatch_to_redliner
)

@pytest.fixture
def sample_docx_file(tmp_path) -> Path:
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    doc.add_paragraph("This Master Services Agreement is entered into between Customer and Vendor.")
    doc.add_heading("Limitation of Liability", level=2)
    doc.add_paragraph("Vendor liability shall be uncapped for any operational losses.")
    doc.add_heading("Governing Law", level=2)
    doc.add_paragraph("This Agreement shall be governed by Delaware law.")
    
    file_path = tmp_path / "sample_contract.docx"
    doc.save(str(file_path))
    return file_path

def test_state_machine_valid_transitions():
    ctx = SessionContext(contract_name="test_contract")
    sm = AuditStateMachine(ctx)
    
    assert ctx.current_state == SessionStateEnum.UNINITIALIZED
    assert sm.can_transition_to(SessionStateEnum.INGESTED) is True
    
    sm.transition_to(SessionStateEnum.INGESTED)
    assert ctx.current_state == SessionStateEnum.INGESTED
    
    sm.transition_to(SessionStateEnum.AUDITING)
    assert ctx.current_state == SessionStateEnum.AUDITING

    sm.transition_to(SessionStateEnum.CRITIQUED)
    assert ctx.current_state == SessionStateEnum.CRITIQUED

    sm.transition_to(SessionStateEnum.REDLINING)
    assert ctx.current_state == SessionStateEnum.REDLINING

    sm.transition_to(SessionStateEnum.AWAITING_HUMAN)
    assert ctx.current_state == SessionStateEnum.AWAITING_HUMAN

    sm.transition_to(SessionStateEnum.FINALIZED)
    assert ctx.current_state == SessionStateEnum.FINALIZED
    assert len(ctx.history) == 6

def test_state_machine_invalid_transition_raises_error():
    ctx = SessionContext()
    sm = AuditStateMachine(ctx)
    
    with pytest.raises(ValueError, match="Invalid state transition"):
        sm.transition_to(SessionStateEnum.FINALIZED)

def test_state_machine_failed_recovery_flow():
    ctx = SessionContext()
    sm = AuditStateMachine(ctx)
    
    sm.transition_to(SessionStateEnum.INGESTED)
    sm.transition_to(SessionStateEnum.FAILED, metadata={"error": "Parsing failed"})
    assert ctx.current_state == SessionStateEnum.FAILED
    
    # Recovery from FAILED back to UNINITIALIZED
    sm.transition_to(SessionStateEnum.UNINITIALIZED)
    assert ctx.current_state == SessionStateEnum.UNINITIALIZED

def test_orchestrator_tools_delegation(sample_docx_file):
    # Test ingestion tool
    chunks = dispatch_to_ingestion(str(sample_docx_file))
    assert len(chunks) > 0
    
    # Test auditor tool
    audit_report = dispatch_to_auditor(chunks, playbook_name="sample_vendor_msa", contract_name="test_doc")
    assert audit_report.total_clauses == len(chunks)
    
    # Test critic tool
    critic_report = dispatch_to_critic(audit_report, playbook_name="sample_vendor_msa")
    assert critic_report.total_verdicts_checked == audit_report.total_clauses
    
    # Test redliner tool
    redline_pkg = dispatch_to_redliner(audit_report, chunks)
    assert redline_pkg.contract_name == "test_doc"

def test_orchestrator_full_pipeline_run(sample_docx_file):
    agent = OrchestratorAgent()
    session = agent.run_audit_pipeline(
        contract_path=str(sample_docx_file),
        playbook_name="sample_vendor_msa",
        session_id="test_session_101"
    )

    assert session.session_id == "test_session_101"
    assert session.current_state == SessionStateEnum.AWAITING_HUMAN
    assert len(session.chunks) > 0
    assert session.audit_report is not None
    assert session.critic_report is not None
    assert session.redline_package is not None

def test_orchestrator_human_review_and_finalization(sample_docx_file, tmp_path):
    agent = OrchestratorAgent()
    session = agent.run_audit_pipeline(
        contract_path=str(sample_docx_file),
        playbook_name="sample_vendor_msa"
    )

    assert session.current_state == SessionStateEnum.AWAITING_HUMAN
    initial_edits_count = len(session.redline_package.edits)

    if initial_edits_count > 0:
        target_clause_id = session.redline_package.edits[0].clause_id
        
        # Test applying human decision: EDIT
        human_decisions = [
            {
                "clause_id": target_clause_id,
                "action": "EDIT",
                "custom_text": "Vendor total liability under this Agreement shall not exceed $500,000."
            }
        ]
        
        session = agent.apply_human_review(session, human_decisions)
        assert target_clause_id in session.human_decisions
        assert session.human_decisions[target_clause_id].action == HumanDecisionEnum.EDIT

    # Finalize document
    out_docx = tmp_path / "final_contract.docx"
    final_path = agent.finalize_review(session, output_path=str(out_docx))
    
    assert session.current_state == SessionStateEnum.FINALIZED
    assert session.final_docx_path == str(out_docx)
    assert Path(final_path).exists()

def test_orchestrator_human_review_rejection_removes_edits(sample_docx_file):
    agent = OrchestratorAgent()
    session = agent.run_audit_pipeline(
        contract_path=str(sample_docx_file),
        playbook_name="sample_vendor_msa"
    )
    
    if session.redline_package and len(session.redline_package.edits) > 0:
        initial_count = len(session.redline_package.edits)
        target_clause_id = session.redline_package.edits[0].clause_id
        
        # Human rejects the proposed edit
        decisions = [{"clause_id": target_clause_id, "action": "REJECT"}]
        session = agent.apply_human_review(session, decisions)
        
        assert len(session.redline_package.edits) == initial_count - 1
        assert not any(e.clause_id == target_clause_id for e in session.redline_package.edits)

def test_orchestrator_invalid_state_operations(sample_docx_file):
    agent = OrchestratorAgent()
    session = SessionContext(contract_path=str(sample_docx_file))
    
    # State is UNINITIALIZED, applying human review or finalization should raise ValueError
    with pytest.raises(ValueError, match="Cannot apply human review in state"):
        agent.apply_human_review(session, [])
        
    with pytest.raises(ValueError, match="Cannot finalize session in state"):
        agent.finalize_review(session)

def test_session_context_serialization(sample_docx_file):
    agent = OrchestratorAgent()
    session = agent.run_audit_pipeline(contract_path=str(sample_docx_file))
    
    dump = session.model_dump()
    assert dump["session_id"] == session.session_id
    assert dump["current_state"] == SessionStateEnum.AWAITING_HUMAN.value
    
    reconstituted = SessionContext.model_validate(dump)
    assert reconstituted.session_id == session.session_id
    assert reconstituted.current_state == SessionStateEnum.AWAITING_HUMAN
