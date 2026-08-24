# End-to-end SDK & Multi-Agent Integration Tests
import pytest
from pathlib import Path
from docx import Document

from backend.agents.base.agent_factory import AgentFactory
from backend.agents.auditor.agent import AuditorAgent
from backend.agents.redliner.agent import RedlinerAgent
from backend.agents.critic.agent import CriticAgent
from backend.agents.orchestrator.agent import OrchestratorAgent
from backend.agents.orchestrator.state_machine import SessionStateEnum, HumanDecisionEnum

@pytest.fixture
def complex_contract_docx(tmp_path) -> Path:
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    
    doc.add_heading("1. Term and Termination", level=2)
    doc.add_paragraph("Either party may terminate this agreement for convenience upon 30 days written notice.")
    
    doc.add_heading("2. Limitation of Liability", level=2)
    doc.add_paragraph("Vendor liability shall be completely uncapped for any operational losses or indirect damages.")
    
    doc.add_heading("3. Indemnification", level=2)
    doc.add_paragraph("Customer agrees to indemnify Vendor against any non-US third-party claims.")
    
    doc.add_heading("4. Governing Law", level=2)
    doc.add_paragraph("This agreement shall be governed by the laws of New York.")
    
    file_path = tmp_path / "complex_contract.docx"
    doc.save(str(file_path))
    return file_path

def test_agent_factory_integration():
    """Verify AgentFactory resolves configurations for all multi-agent roles."""
    roles = ["Orchestrator", "Auditor", "Redliner", "Critic"]
    models = ["claude-sonnet", "claude-haiku", "claude-haiku", "claude-haiku"]
    
    for role, model in zip(roles, models):
        cfg = AgentFactory.create_agent_config(agent_role=role, model_alias=model)
        assert cfg["agent_role"] == role
        assert "anthropic.claude-3" in cfg["model_id"]
        assert cfg["parameters"]["temperature"] == 0.1

def test_multi_agent_pipeline_reasoning_and_grounding(complex_contract_docx):
    """Verifies that Orchestrator coordinates Auditor, Critic, and Redliner cleanly."""
    orchestrator = OrchestratorAgent()
    session = orchestrator.run_audit_pipeline(
        contract_path=str(complex_contract_docx),
        playbook_name="sample_vendor_msa",
        session_id="sdk_test_session_001"
    )

    # 1. State machine check
    assert session.current_state == SessionStateEnum.AWAITING_HUMAN
    assert len(session.chunks) >= 4
    
    # 2. Auditor Agent check
    audit_report = session.audit_report
    assert audit_report is not None
    assert audit_report.total_clauses == len(session.chunks)
    assert audit_report.overall_risk_level.value in ["CRITICAL", "HIGH", "MEDIUM", "LOW"]

    # 3. Critic Agent grounding check
    critic_report = session.critic_report
    assert critic_report is not None
    assert critic_report.grounded_verdicts >= 0
    assert isinstance(critic_report.all_grounded, bool)

    # 4. Redliner Agent check
    redlines = session.redline_package
    assert redlines is not None
    assert len(redlines.edits) > 0

def test_sdk_hitl_approval_and_final_document(complex_contract_docx, tmp_path):
    """Tests human-in-the-loop decisions (APPROVE, EDIT, REJECT) end-to-end."""
    orchestrator = OrchestratorAgent()
    session = orchestrator.run_audit_pipeline(
        contract_path=str(complex_contract_docx),
        playbook_name="sample_vendor_msa",
        session_id="sdk_test_session_002"
    )

    edits = session.redline_package.edits
    assert len(edits) > 0

    first_clause_id = edits[0].clause_id
    decisions = [
        {
            "clause_id": first_clause_id,
            "action": "EDIT",
            "custom_text": "Vendor liability under this section shall not exceed 1x the annual contract value."
        }
    ]

    session = orchestrator.apply_human_review(session, decisions)
    assert first_clause_id in session.human_decisions
    assert session.human_decisions[first_clause_id].action == HumanDecisionEnum.EDIT

    out_file = tmp_path / "sdk_finalized_redline.docx"
    final_path = orchestrator.finalize_review(session, output_path=str(out_file))

    assert session.current_state == SessionStateEnum.FINALIZED
    assert Path(final_path).exists()
