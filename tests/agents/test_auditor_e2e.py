# End-to-end integration test for contract document auditing against legal playbooks.
from pathlib import Path
from docx import Document
from backend.ingestion.pipeline import ingest_contract
from backend.agents.auditor.agent import AuditorAgent
from backend.agents.auditor.verdict_schema import VerdictEnum, SeverityEnum, ContractAuditReport

def test_e2e_contract_auditing(tmp_path: Path):
    docx_path = tmp_path / "audit_test_contract.docx"
    doc = Document()
    doc.add_heading("Master Services Agreement", level=1)
    
    # Clause 1: Compliant governing law
    doc.add_heading("Section 1: Governing Law", level=2)
    doc.add_paragraph("This Agreement shall be governed by Delaware law.")
    
    # Clause 2: Severe deviation (uncapped liability)
    doc.add_heading("Section 2: Limitation of Liability", level=2)
    doc.add_paragraph("Vendor liability shall be uncapped for any operational losses or property damage.")
    
    # Clause 3: Fallback position (2x fees cap)
    doc.add_heading("Section 3: Alternate Cap", level=2)
    doc.add_paragraph("Total liability shall be capped at 2x fees paid in preceding 12 months.")
    
    doc.save(str(docx_path))
    
    # Run ingestion pipeline
    chunks = ingest_contract(docx_path)
    assert len(chunks) == 3
    
    # Run auditor agent
    auditor = AuditorAgent()
    report = auditor.audit_contract(chunks, playbook_name="sample_vendor_msa", contract_name="audit_test_contract.docx")
    
    assert report.total_clauses == 3
    assert report.overall_risk_level == SeverityEnum.CRITICAL
    assert report.verdicts[0].verdict == VerdictEnum.COMPLIANT
    assert report.verdicts[1].verdict == VerdictEnum.DEVIATION
    assert report.verdicts[1].severity == SeverityEnum.CRITICAL
    assert report.verdicts[2].verdict == VerdictEnum.DEVIATION
    assert report.verdicts[2].severity == SeverityEnum.MEDIUM


def test_e2e_auditing_indemnification_violation(tmp_path: Path):
    docx_path = tmp_path / "indemnity_contract.docx"
    doc = Document()
    doc.add_heading("Vendor Agreement", level=1)
    doc.add_heading("Section 1: Indemnification", level=2)
    doc.add_paragraph("Customer indemnifying Vendor for any general use of Vendor's software or services.")
    doc.save(str(docx_path))

    chunks = ingest_contract(docx_path)
    auditor = AuditorAgent()
    report = auditor.audit_contract(chunks, playbook_name="sample_vendor_msa", contract_name="indemnity_contract.docx")

    assert report.total_clauses == 1
    assert report.overall_risk_level == SeverityEnum.CRITICAL
    v = report.verdicts[0]
    assert v.verdict == VerdictEnum.DEVIATION
    assert v.severity == SeverityEnum.CRITICAL
    assert len(v.playbook_citation_ids) > 0


def test_e2e_auditing_jurisdiction_fallback(tmp_path: Path):
    docx_path = tmp_path / "jurisdiction_contract.docx"
    doc = Document()
    doc.add_heading("Vendor Agreement", level=1)
    doc.add_heading("Section 1: Governing Law", level=2)
    doc.add_paragraph("This agreement is governed by the laws of New York.")
    doc.save(str(docx_path))

    chunks = ingest_contract(docx_path)
    auditor = AuditorAgent()
    report = auditor.audit_contract(chunks, playbook_name="sample_vendor_msa", contract_name="jurisdiction_contract.docx")

    assert report.total_clauses == 1
    v = report.verdicts[0]
    assert v.verdict == VerdictEnum.DEVIATION
    assert v.severity == SeverityEnum.MEDIUM


def test_e2e_report_json_serialization(tmp_path: Path):
    docx_path = tmp_path / "simple_contract.docx"
    doc = Document()
    doc.add_heading("Simple Contract", level=1)
    doc.add_paragraph("Governing law shall be Delaware law.")
    doc.save(str(docx_path))

    chunks = ingest_contract(docx_path)
    auditor = AuditorAgent()
    report = auditor.audit_contract(chunks, playbook_name="sample_vendor_msa", contract_name="simple_contract.docx")

    dump = report.model_dump()
    assert "contract_name" in dump
    assert "playbook_name" in dump
    assert "total_clauses" in dump
    assert "verdicts" in dump
    assert "overall_risk_level" in dump
